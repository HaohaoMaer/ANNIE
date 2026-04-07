"""DOCX Reader Tool - Reads and extracts content from DOCX files.

Used for reading game flow and background documents.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document

from annie.npc.tools.base_tool import BaseTool


class DOCXReaderTool(BaseTool):
    """Tool for reading DOCX documents."""

    name = "docx_reader"
    description = (
        "Reads and extracts text content from DOCX documents. "
        "Can extract full text or structured content by paragraphs. "
        "Use this tool when you need to read game flow or background documents."
    )

    def execute(self, context: dict) -> dict:
        """Execute DOCX reading operation.

        Args:
            context: Dict with keys:
                - 'task': str, the task description
                - 'npc_name': str, the NPC's name
                - 'docx_path': str, path to DOCX file
                - 'operation': str (optional), 'full' or 'structured' (default: 'full')

        Returns:
            Dict with:
                - 'success': bool
                - 'text': str, extracted text
                - 'paragraphs': list, list of paragraphs (if structured)
                - 'paragraph_count': int, total paragraphs
                - 'error': str (if failed)
        """
        docx_path = context.get("docx_path")
        operation = context.get("operation", "full")

        if not docx_path:
            return self._error_result("No DOCX path provided")

        path = Path(docx_path)
        if not path.exists():
            return self._error_result(f"DOCX file not found: {docx_path}")

        if path.suffix.lower() != ".docx":
            return self._error_result(f"Not a DOCX file: {path.suffix}")

        try:
            doc = Document(str(path))

            if operation == "structured":
                return self._extract_structured(doc, path)
            else:
                return self._extract_full(doc, path)

        except Exception as e:
            return self._error_result(f"Failed to read DOCX: {str(e)}")

    def _extract_full(self, doc: Document, path: Path) -> dict[str, Any]:
        """Extract full text from document."""
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        full_text = "\n\n".join(paragraphs)

        return {
            "success": True,
            "text": full_text,
            "paragraph_count": len(paragraphs),
            "file_name": path.name,
        }

    def _extract_structured(self, doc: Document, path: Path) -> dict[str, Any]:
        """Extract structured content from document."""
        paragraphs = []
        current_section = "正文"
        section_content = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            is_heading = False
            if para.style and "heading" in para.style.name.lower():
                is_heading = True

            if is_heading and section_content:
                paragraphs.append({
                    "section": current_section,
                    "content": "\n".join(section_content),
                })
                current_section = text
                section_content = []
            else:
                section_content.append(text)

        if section_content:
            paragraphs.append({
                "section": current_section,
                "content": "\n".join(section_content),
            })

        full_text = "\n\n".join(
            f"【{p['section']}】\n{p['content']}"
            for p in paragraphs
        )

        return {
            "success": True,
            "text": full_text,
            "paragraphs": paragraphs,
            "paragraph_count": len(paragraphs),
            "file_name": path.name,
        }

    def _error_result(self, error: str) -> dict[str, Any]:
        """Create an error result dict."""
        return {
            "success": False,
            "text": "",
            "paragraphs": [],
            "paragraph_count": 0,
            "error": error,
        }

    def get_document_info(self, docx_path: str) -> dict[str, Any]:
        """Get basic information about a DOCX file.

        Args:
            docx_path: Path to the DOCX file.

        Returns:
            Dict with document metadata.
        """
        path = Path(docx_path)
        if not path.exists():
            return {"error": "File not found"}

        try:
            doc = Document(str(path))
            return {
                "file_name": path.name,
                "file_size": path.stat().st_size,
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "section_count": len(doc.sections),
            }
        except Exception as e:
            return {"error": str(e)}

    def extract_tables(self, docx_path: str) -> list[list[list[str]]]:
        """Extract tables from a DOCX file.

        Args:
            docx_path: Path to the DOCX file.

        Returns:
            List of tables, each table is a list of rows,
            each row is a list of cell texts.
        """
        path = Path(docx_path)
        if not path.exists():
            return []

        try:
            doc = Document(str(path))
            tables = []

            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)

            return tables

        except Exception:
            return []
